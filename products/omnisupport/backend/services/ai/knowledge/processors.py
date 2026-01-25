"""Document processors for text extraction."""

import logging
import re
from abc import ABC, abstractmethod
from pathlib import Path

logger = logging.getLogger(__name__)


class BaseProcessor(ABC):
    """Base document processor."""

    @abstractmethod
    def extract_text(self, content: bytes, filename: str = "") -> str | None:
        """Extract text from document content."""
        pass

    def clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)
        # Remove control characters
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
        return text.strip()


class PDFProcessor(BaseProcessor):
    """PDF document processor using pypdf."""

    def extract_text(self, content: bytes, filename: str = "") -> str | None:
        """Extract text from PDF."""
        try:
            from pypdf import PdfReader
            import io

            reader = PdfReader(io.BytesIO(content))
            text_parts = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            if text_parts:
                return self.clean_text("\n\n".join(text_parts))

        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")

        return None


class DOCXProcessor(BaseProcessor):
    """DOCX document processor using python-docx."""

    def extract_text(self, content: bytes, filename: str = "") -> str | None:
        """Extract text from DOCX."""
        try:
            from docx import Document
            import io

            doc = Document(io.BytesIO(content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            if text_parts:
                return self.clean_text("\n\n".join(text_parts))

        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")

        return None


class TXTProcessor(BaseProcessor):
    """Plain text processor."""

    def extract_text(self, content: bytes, filename: str = "") -> str | None:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            for encoding in ["utf-8", "cp1251", "latin-1"]:
                try:
                    text = content.decode(encoding)
                    return self.clean_text(text)
                except UnicodeDecodeError:
                    continue

        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")

        return None


class HTMLProcessor(BaseProcessor):
    """HTML document processor using BeautifulSoup."""

    def extract_text(self, content: bytes, filename: str = "") -> str | None:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup

            # Decode content
            text_content = None
            for encoding in ["utf-8", "cp1251", "latin-1"]:
                try:
                    text_content = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if not text_content:
                return None

            soup = BeautifulSoup(text_content, "lxml")

            # Remove unwanted elements
            for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
                tag.decompose()

            # Get text
            text = soup.get_text(separator="\n", strip=True)

            if text:
                return self.clean_text(text)

        except Exception as e:
            logger.error(f"Error extracting HTML text: {e}")

        return None


class MarkdownProcessor(BaseProcessor):
    """Markdown processor."""

    def extract_text(self, content: bytes, filename: str = "") -> str | None:
        """Extract text from Markdown, preserving structure."""
        try:
            text = content.decode("utf-8")

            # Remove code blocks but keep inline code
            text = re.sub(r"```[\s\S]*?```", "[code block]", text)

            # Remove images
            text = re.sub(r"!\[.*?\]\(.*?\)", "", text)

            # Convert links to text
            text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)

            # Remove HTML tags
            text = re.sub(r"<[^>]+>", "", text)

            return self.clean_text(text)

        except Exception as e:
            logger.error(f"Error extracting Markdown text: {e}")

        return None


class DocumentProcessorFactory:
    """Factory for document processors."""

    _processors = {
        "application/pdf": PDFProcessor,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DOCXProcessor,
        "text/plain": TXTProcessor,
        "text/html": HTMLProcessor,
        "text/markdown": MarkdownProcessor,
    }

    _extension_map = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".html": "text/html",
        ".htm": "text/html",
        ".md": "text/markdown",
        ".markdown": "text/markdown",
    }

    @classmethod
    def get_processor(
        cls,
        mime_type: str | None = None,
        filename: str | None = None,
    ) -> BaseProcessor | None:
        """Get appropriate processor for document type."""
        # Determine mime type from extension if not provided
        if not mime_type and filename:
            ext = Path(filename).suffix.lower()
            mime_type = cls._extension_map.get(ext)

        if mime_type:
            processor_class = cls._processors.get(mime_type)
            if processor_class:
                return processor_class()

        return None

    @classmethod
    def extract_text(
        cls,
        content: bytes,
        mime_type: str | None = None,
        filename: str | None = None,
    ) -> str | None:
        """Extract text using appropriate processor."""
        processor = cls.get_processor(mime_type, filename)
        if processor:
            return processor.extract_text(content, filename or "")
        return None

    @classmethod
    def supported_types(cls) -> list[str]:
        """Get list of supported MIME types."""
        return list(cls._processors.keys())

    @classmethod
    def supported_extensions(cls) -> list[str]:
        """Get list of supported file extensions."""
        return list(cls._extension_map.keys())
